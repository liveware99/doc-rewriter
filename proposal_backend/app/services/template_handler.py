
from docx import Document

def load_template(path):
    return Document(path)

def extract_placeholders(document):
    placeholders = []
    for para in document.paragraphs:
        if "<" in para.text and ">" in para.text:
            placeholders.append(para.text.strip())
    return placeholders

def replace_placeholder(document, placeholder, new_text):
    for para in document.paragraphs:
        if placeholder in para.text:
            para.text = para.text.replace(placeholder, new_text)
